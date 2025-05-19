#GIUSTO
import docx
from docx.oxml.ns import nsmap

def replace_text_in_docx(input_file, output_file, replacements):
    doc = docx.Document(input_file)
    
    def replace_in_paragraph(paragraph):
        modified = False
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                inline = paragraph._p.xpath('.//w:r')
                # Collezione di tutti i run che contengono parti del testo da sostituire
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        modified = True
        return modified
    
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    for section in doc.sections:
        for header in [section.header]:
            for paragraph in header.paragraphs:
                replace_in_paragraph(paragraph)
        
        for footer in [section.footer]:
            for paragraph in footer.paragraphs:
                replace_in_paragraph(paragraph)
    
    doc.save(output_file)

if __name__ == "__main__":
    sostituzioni = {
        "[[CITTA_SEDE]]": "Modena",
        "[[VIA]]": "perfavore funziona, 2/Q",
        "[[NOME_RAP]]": "Mario Rossi",
        "[[AZIENDA]]": "Tecnologia",
        "[[STUDENTE]]": "Chiara Bianchi",
        "[[INIZIO_PERIODO]]": "12/12/2025",
        "[[FINE_PERIODO]]": "1/1/2026"
    }

    replace_text_in_docx(
        "template.docx", 
        f"{AZIENDA}_convenzioni.docx", 
        sostituzioni
    )